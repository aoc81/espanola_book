from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[2]
TITLE_FILE = "Front Matter 00 Title and Edition.md"
BODY_FILES = [
    "Front Matter 01 Author Note.md",
    "Front Matter 03 Note on AI Assistance.md",
    "Prologue.md",
    "Chapter 01.md",
    "Chapter 02.md",
    "Chapter 03.md",
    "Chapter 04.md",
    "Chapter 05.md",
    "Chapter 06.md",
    "Chapter 07.md",
    "Chapter 08.md",
    "Chapter 09.md",
    "Chapter 10.md",
    "Chapter 11.md",
    "Chapter 12.md",
    "Chapter 13.md",
    "Chapter 14.md",
    "Chapter 15.md",
    "Chapter 16.md",
    "Chapter 17.md",
    "Chapter 18.md",
    "Epilogue.md",
    "Front Matter 02 Methodology and Limitations.md",
    "Front Matter 04 Note on Naming and Terminology.md",
    "Appendix Source Notes.md",
    "Appendix Reference Links Guide.md",
]
DEFAULT_OUTPUTS = {
    ("v4", "en"): "From Hooligans to War Machines - v4 Book.docx",
    ("v4.1", "en"): "From Hooligans to War Machines - v4.1 Book.docx",
    ("v4.1", "es"): "De hooligans a máquinas de guerra - v4.1-es.docx",
    ("v4.1", "fr"): "Des hooligans aux machines de guerre - v4.1-fr.docx",
    ("v4.1", "ru"): "Espanola - ot huliganov do boevyh mashin - v4.1-ru.docx",
}
DEFAULT_FALLBACK = "book-output.updated.docx"
DEFAULT_COVER = REPO_ROOT / "assets" / "images" / "covers" / "Espanola Book Cover-New.png"
INLINE_TOKEN_RE = re.compile(r"\[\^([^\]]+)\]|\[([^\]]+)\]\((https?://[^)]+)\)|(https?://\S+)")
EMPHASIS_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


@dataclass
class MarkdownDocument:
    title: str
    blocks: list[dict]
    footnotes: dict[str, str]


@dataclass
class EditionPaths:
    version: str
    language: str
    contents_dir: Path
    manifest_path: Path
    image_root: Path
    output_dir: Path
    cover_path: Path | None
    output_name: str


class FootnoteManager:
    def __init__(self) -> None:
        self._label_to_number: dict[str, int] = {}
        self._notes: list[tuple[int, str]] = []

    @property
    def notes(self) -> list[tuple[int, str]]:
        return self._notes

    def register(self, label: str, text: str) -> int:
        if label in self._label_to_number:
            return self._label_to_number[label]
        number = len(self._notes) + 1
        self._label_to_number[label] = number
        self._notes.append((number, text))
        return number


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig").replace("\r\n", "\n").replace("\r", "\n")


def strip_markdown_heading(line: str) -> str:
    stripped = line.strip()
    if stripped.startswith("#"):
        return stripped.lstrip("#").strip()
    return stripped


def parse_markdown(text: str) -> MarkdownDocument:
    lines = [line.rstrip() for line in text.split("\n")]
    index = 0
    while index < len(lines) and not lines[index].strip():
        index += 1
    if index == len(lines):
        raise ValueError("Source file is empty.")

    title = strip_markdown_heading(lines[index])
    footnotes: dict[str, str] = {}
    content_lines: list[str] = []
    footnote_pattern = re.compile(r"^\[\^([^\]]+)\]:\s*(.*)$")

    for line in lines[index + 1 :]:
        match = footnote_pattern.match(line)
        if match:
            footnotes[match.group(1)] = match.group(2).strip()
            continue
        content_lines.append(line)

    blocks: list[dict] = []
    current: list[str] = []

    def flush_current() -> None:
        nonlocal current
        if not current:
            return
        if len(current) >= 2 and all(line.strip().startswith("|") and line.strip().endswith("|") for line in current):
            rows = []
            for raw_line in current:
                cells = [cell.strip() for cell in raw_line.strip()[1:-1].split("|")]
                rows.append(cells)
            if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
                blocks.append({"type": "table", "headers": rows[0], "rows": rows[2:]})
                current = []
                return
        if all(item.lstrip().startswith("- ") for item in current):
            blocks.append({"type": "list", "items": [item.lstrip()[2:].strip() for item in current]})
        else:
            blocks.append({"type": "paragraph", "lines": current})
        current = []

    for line in content_lines:
        stripped = line.strip()
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            continue
        if stripped == "---":
            flush_current()
            continue
        if stripped.startswith("## "):
            flush_current()
            blocks.append({"type": "subheading", "text": stripped[3:].strip()})
            continue
        if stripped:
            current.append(line)
            continue
        flush_current()
    flush_current()

    return MarkdownDocument(title=title, blocks=blocks, footnotes=footnotes)


def title_lines(text: str) -> list[str]:
    return [strip_markdown_heading(line) for line in text.split("\n") if line.strip()]


def ensure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.3)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title_style = styles["Title"]
    title_style.font.name = "Times New Roman"
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(10)

    subtitle_style = styles["Subtitle"]
    subtitle_style.font.name = "Times New Roman"
    subtitle_style.font.size = Pt(12)
    subtitle_style.font.italic = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_before = Pt(0 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(10 if style_name == "Heading 1" else 4)
        style.paragraph_format.keep_with_next = True

    custom_styles = {
        "Book Center": {"base": "Normal", "size": 12, "align": WD_ALIGN_PARAGRAPH.CENTER, "italic": False},
        "Book Bullet": {"base": "Normal", "size": 12, "align": WD_ALIGN_PARAGRAPH.JUSTIFY, "italic": False},
        "Contents Title": {"base": "Normal", "size": 16, "align": WD_ALIGN_PARAGRAPH.LEFT, "italic": False},
        "Image Caption": {"base": "Normal", "size": 10, "align": WD_ALIGN_PARAGRAPH.CENTER, "italic": True},
        "Image Source": {"base": "Normal", "size": 9, "align": WD_ALIGN_PARAGRAPH.CENTER, "italic": False},
        "Chapter References": {"base": "Normal", "size": 8, "align": WD_ALIGN_PARAGRAPH.LEFT, "italic": False},
        "Appendix Body": {"base": "Normal", "size": 10, "align": WD_ALIGN_PARAGRAPH.LEFT, "italic": False},
        "Appendix Bullet": {"base": "Normal", "size": 10, "align": WD_ALIGN_PARAGRAPH.LEFT, "italic": False},
    }

    for style_name, spec in custom_styles.items():
        if style_name in styles:
            continue
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[spec["base"]]
        style.font.name = "Times New Roman"
        style.font.size = Pt(spec["size"])
        style.font.italic = spec["italic"]
        style.paragraph_format.alignment = spec["align"]
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.space_after = Pt(8 if style_name == "Book Center" else 4 if style_name == "Book Bullet" else 10 if style_name == "Contents Title" else 3)

    styles["Book Bullet"].paragraph_format.left_indent = Inches(0.25)
    styles["Appendix Body"].paragraph_format.first_line_indent = Inches(0)
    styles["Appendix Body"].paragraph_format.space_after = Pt(6)
    styles["Appendix Body"].paragraph_format.line_spacing = 1.1
    styles["Appendix Bullet"].paragraph_format.left_indent = Inches(0.25)
    styles["Appendix Bullet"].paragraph_format.first_line_indent = Inches(0)
    styles["Appendix Bullet"].paragraph_format.space_after = Pt(4)
    styles["Image Source"].paragraph_format.space_after = Pt(10)
    styles["Chapter References"].paragraph_format.line_spacing = 1.0


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def add_page_number(paragraph) -> None:
    add_field(paragraph, "PAGE")


def enable_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_section_layout(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def configure_cover_layout(section) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def set_page_number_format(section, *, fmt: str | None = None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    if fmt is not None:
        pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))


def setup_footer_with_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(paragraph)


def add_toc(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    paragraph.add_run()._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-1" \\h \\z \\u'
    paragraph.add_run()._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    paragraph.add_run()._r.append(separate)
    paragraph.add_run("Update field in Word if the contents does not appear automatically.")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(end)


def add_hyperlink(paragraph, text: str, url: str, *, size: Pt | None = None) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    if size is not None:
        size_node = OxmlElement("w:sz")
        size_node.set(qn("w:val"), str(int(size.pt * 2)))
        rpr.append(size_node)
        size_cs_node = OxmlElement("w:szCs")
        size_cs_node.set(qn("w:val"), str(int(size.pt * 2)))
        rpr.append(size_cs_node)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(rpr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_footnote_reference(paragraph, number: int) -> None:
    run = paragraph.add_run(str(number))
    run.font.superscript = True
    run.font.size = Pt(8)


def add_text_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: Pt | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = size


def add_emphasis_runs(paragraph, text: str, *, size: Pt | None = None) -> None:
    position = 0
    for match in EMPHASIS_RE.finditer(text):
        if match.start() > position:
            add_text_run(paragraph, text[position:match.start()], size=size)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            add_text_run(paragraph, token[2:-2], bold=True, size=size)
        elif token.startswith("*") and token.endswith("*"):
            add_text_run(paragraph, token[1:-1], italic=True, size=size)
        elif token.startswith("`") and token.endswith("`"):
            add_text_run(paragraph, token[1:-1], size=size)
        position = match.end()
    if position < len(text):
        add_text_run(paragraph, text[position:], size=size)


def add_inline_markdown(
    paragraph,
    text: str,
    footnotes: dict[str, str],
    footnote_manager: FootnoteManager,
    *,
    text_size: Pt | None = None,
) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            add_emphasis_runs(paragraph, text[position : match.start()], size=text_size)

        footnote_label, link_text, link_url, bare_url = match.groups()
        if footnote_label:
            note_text = footnotes.get(footnote_label)
            if note_text is None:
                note_text = f"SOURCE GAP: missing footnote definition for {footnote_label}."
            number = footnote_manager.register(footnote_label, note_text)
            add_footnote_reference(paragraph, number)
        elif link_text and link_url:
            add_hyperlink(paragraph, link_text, link_url, size=text_size)
        elif bare_url:
            clean_url = bare_url.rstrip(".,;")
            add_hyperlink(paragraph, clean_url, clean_url, size=text_size)
            trailing = bare_url[len(clean_url) :]
            if trailing:
                add_emphasis_runs(paragraph, trailing, size=text_size)

        position = match.end()

    if position < len(text):
        add_emphasis_runs(paragraph, text[position:], size=text_size)


def add_title_page(document: Document, lines: list[str]) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(140)
    title = document.add_paragraph(lines[0], style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in lines[1:]:
        para = document.add_paragraph(line, style="Subtitle")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover_page(document: Document, cover_path: Path) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    configure_cover_layout(section)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(cover_path), width=section.page_width, height=section.page_height)


def add_contents_page(document: Document) -> None:
    document.add_paragraph("Contents", style="Contents Title")
    toc_para = document.add_paragraph(style="Normal")
    toc_para.paragraph_format.first_line_indent = Inches(0)
    add_toc(toc_para)


def add_markdown_lines_as_paragraph(
    paragraph,
    lines: list[str],
    footnotes: dict[str, str],
    footnote_manager: FootnoteManager,
    *,
    text_size: Pt | None = None,
) -> None:
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        add_inline_markdown(paragraph, line, footnotes, footnote_manager, text_size=text_size)


def add_markdown_table(document: Document, headers: list[str], rows: list[list[str]], *, text_size: Pt | None = None) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    for col, header in enumerate(headers):
        paragraph = table.rows[0].cells[col].paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        if text_size is not None:
            run.font.size = text_size

    for row_index, row in enumerate(rows, start=1):
        for col, value in enumerate(row):
            paragraph = table.rows[row_index].cells[col].paragraphs[0]
            add_emphasis_runs(paragraph, value, size=text_size)


def load_image_manifest(manifest_path: Path) -> dict[str, list[dict]]:
    if not manifest_path.exists():
        return {}
    entries = json.loads(manifest_path.read_text(encoding="utf-8"))
    grouped: dict[str, list[dict]] = {}
    for entry in entries:
        grouped.setdefault(entry["chapter_file"], []).append(entry)
    for chapter_entries in grouped.values():
        chapter_entries.sort(key=lambda entry: entry.get("relative_path", ""))
    return grouped


def add_chapter_images(document: Document, image_root: Path, chapter_file: str, image_manifest: dict[str, list[dict]]) -> None:
    entries = image_manifest.get(chapter_file, [])
    if not entries:
        return

    document.add_paragraph("", style="Normal")
    document.add_heading("Selected Images", level=2)

    for entry in entries:
        image_path = image_root / entry["relative_path"]
        if not image_path.exists():
            continue

        picture_para = document.add_paragraph(style="Book Center")
        picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_para.add_run().add_picture(str(image_path), width=Inches(5.8))

        caption_para = document.add_paragraph(entry["caption"], style="Image Caption")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        source_para = document.add_paragraph(style="Image Source")
        source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text_run(source_para, "Source: ", bold=True)
        add_text_run(source_para, f"{entry['source_name']}. ")
        add_text_run(source_para, "Author: ", bold=True)
        add_text_run(source_para, f"{entry['author']}. ")
        add_text_run(source_para, "License: ", bold=True)
        add_text_run(source_para, f"{entry['license']}. ")
        add_text_run(source_para, "Original: ", bold=True)
        if entry.get("media_url"):
            add_hyperlink(source_para, entry["media_url"], entry["media_url"])
        elif entry.get("source_url"):
            add_hyperlink(source_para, entry["source_url"], entry["source_url"])
        else:
            add_text_run(source_para, "N/A")


def add_chapter_references(document: Document, footnote_manager: FootnoteManager) -> None:
    if not footnote_manager.notes:
        return

    document.add_paragraph("", style="Normal")
    document.add_heading("References", level=2)
    for number, text in footnote_manager.notes:
        paragraph = document.add_paragraph(style="Chapter References")
        number_run = paragraph.add_run(f"{number}. ")
        number_run.bold = True
        number_run.font.size = Pt(8)
        text_run = paragraph.add_run(text)
        text_run.font.size = Pt(8)


def add_file(
    document: Document,
    path: Path,
    *,
    break_before: bool = True,
    image_root: Path | None = None,
    image_manifest: dict[str, list[dict]] | None = None,
) -> None:
    footnote_manager = FootnoteManager()
    md_doc = parse_markdown(read_text(path))
    is_appendix = path.name.startswith("Appendix ")
    body_style = "Appendix Body" if is_appendix else "Normal"
    bullet_style = "Appendix Bullet" if is_appendix else "Book Bullet"
    body_text_size = Pt(10) if is_appendix else None
    if break_before:
        document.add_page_break()
    document.add_heading(md_doc.title, level=1)

    for block in md_doc.blocks:
        if block["type"] == "subheading":
            document.add_heading(block["text"], level=2)
            continue
        if block["type"] == "table":
            add_markdown_table(document, block["headers"], block["rows"], text_size=body_text_size)
            continue
        if block["type"] == "list":
            for item in block["items"]:
                para = document.add_paragraph(style=bullet_style)
                add_inline_markdown(para, item, md_doc.footnotes, footnote_manager, text_size=body_text_size)
            continue
        para = document.add_paragraph(style=body_style)
        add_markdown_lines_as_paragraph(para, block["lines"], md_doc.footnotes, footnote_manager, text_size=body_text_size)

    add_chapter_references(document, footnote_manager)

    if image_root is not None and image_manifest is not None:
        add_chapter_images(document, image_root, path.name, image_manifest)


def build_book(paths: EditionPaths, *, include_chapter_images: bool = True) -> Path:
    document = Document()
    ensure_styles(document)
    enable_update_fields_on_open(document)
    image_manifest = load_image_manifest(paths.manifest_path) if include_chapter_images else {}

    has_cover = paths.cover_path is not None and paths.cover_path.exists()

    if has_cover:
        add_cover_page(document, paths.cover_path)
        title_section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        title_section = document.sections[0]

    title_section.different_first_page_header_footer = True
    configure_section_layout(title_section)
    add_title_page(document, title_lines(read_text(paths.contents_dir / TITLE_FILE)))

    front_matter_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(front_matter_section)
    set_page_number_format(front_matter_section, fmt="lowerRoman", start=1)
    setup_footer_with_page_number(front_matter_section)
    add_contents_page(document)

    main_matter_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section_layout(main_matter_section)
    set_page_number_format(main_matter_section, fmt="decimal", start=1)
    setup_footer_with_page_number(main_matter_section)

    for index, filename in enumerate(BODY_FILES):
        add_file(
            document,
            paths.contents_dir / filename,
            break_before=index != 0,
            image_root=paths.image_root if include_chapter_images else None,
            image_manifest=image_manifest if include_chapter_images else None,
        )

    paths.output_dir.mkdir(parents=True, exist_ok=True)
    output_path = paths.output_dir / paths.output_name
    try:
        document.save(output_path)
        return output_path
    except PermissionError:
        fallback_path = paths.output_dir / DEFAULT_FALLBACK
        document.save(fallback_path)
        return fallback_path


def build_paths(version: str, language: str, output_name: str | None, cover_path: str | None, no_cover: bool = False) -> EditionPaths:
    contents_dir = REPO_ROOT / "manuscripts" / version / language / "contents"
    manifest_path = REPO_ROOT / "assets" / "manifests" / version / language / "chapter_images.json"
    image_root = REPO_ROOT / "assets" / "images" / "manuscripts" / version / language
    output_dir = REPO_ROOT / "build" / "docx" / version / language

    if not contents_dir.exists():
        raise FileNotFoundError(f"Edition contents not found: {contents_dir}")

    return EditionPaths(
        version=version,
        language=language,
        contents_dir=contents_dir,
        manifest_path=manifest_path,
        image_root=image_root,
        output_dir=output_dir,
        cover_path=None if no_cover else Path(cover_path) if cover_path else DEFAULT_COVER,
        output_name=output_name or DEFAULT_OUTPUTS.get((version, language), f"book-{version}-{language}.docx"),
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a DOCX book from a canonical manuscript edition.")
    parser.add_argument("--version", default="v4.1", help="Edition version, for example v4 or v4.1.")
    parser.add_argument("--language", default="en", help="Edition language, for example en, es, or fr.")
    parser.add_argument("--output-name", help="Optional DOCX filename override.")
    parser.add_argument("--cover", help="Optional absolute or repo-relative cover image path override.")
    parser.add_argument("--no-cover", action="store_true", help="Omit the cover page from the DOCX.")
    parser.add_argument("--no-chapter-images", action="store_true", help="Omit interior chapter images while keeping the cover.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    cover_arg = args.cover
    if cover_arg and not Path(cover_arg).is_absolute():
        cover_arg = str((REPO_ROOT / cover_arg).resolve())
    paths = build_paths(args.version, args.language, args.output_name, cover_arg, args.no_cover)
    output_path = build_book(paths, include_chapter_images=not args.no_chapter_images)
    print(output_path)


if __name__ == "__main__":
    main()
